import io
import os
import sys
import sqlite3
# import python
#from PIL import Image

from docx import Document
from docx.shared import Inches
from docx.shared import Mm
from docx.shared import Pt


import matplotlib.pyplot as plt

a = list(range(10))
b = list(range(10))
c = list(range(100))


print(b)

plt.plot(a)
plt.scatter(a,b)

buf = io.BytesIO()
#plt.savefig('plot.png')

plt.savefig(buf, format="png")

#img = Image.open(buf)

#img.show()

"""
import io
from PIL import Image
import matplotlib.pyplot as plt

plt.figure()
plt.plot([1, 2])
plt.title("test")
buf = io.BytesIO()
plt.savefig(buf, format='png')
buf.seek(0)
im = Image.open(buf)
im.show()
buf.close()
"""
print(a)


# ---------------------------
document = Document()

document.add_heading('Document Title', 0)

p = document.add_paragraph('A plain paragraph having some ')
p = document.add_paragraph('''Call-off stock refers to the situation where, at the time of the transport of goods to another Member State, the supplier already knows the identity of the person acquiring the goods, to whom these goods will be supplied at a later stage and after they have arrived in the Member State of destination. This currently gives rise to a deemed supply (in the Member State of departure of the goods) and a deemed intra-Community acquisition (in the Member State of arrival of the goods), followed by a ‘domestic’ supply in the Member State of arrival, and requires the supplier to be identified for VAT purposes in that Member State. To avoid this, such transactions, where they take place between two taxable persons should be, under certain conditions, considered to give rise to one exempt supply in the Member State of departure and one intra-Community acquisition in the Member State of arrival.
With the completion of the single market as from the first day of 1993, fiscal controls at intra-Community frontiers were abolished. The elimination of these frontier controls induced a change in fiscal policy towards the acquisition principle.
This policy applied first to the exchanges of goods between taxable persons but was then extended with the exchanges of services between taxable persons.
The basic rule applied to the acquisition principle is that the sale (supply) of goods or services in intra-Community trade in one Member State gives rise to a corresponding acquisition (purchase) in the Member State of arrival of the goods or provision of the services. 
A purchaser may take an intra-Community supply exempt of VAT as long as he meets two simple conditions:''')

p = document.add_paragraph('''Call-off stock refers to the situation where, at the time of the transport of goods to another Member State, the supplier already knows the identity of the person acquiring the goods, to whom these goods will be supplied at a later stage and after they have arrived in the Member State of destination. This currently gives rise to a deemed supply (in the Member State of departure of the goods) and a deemed intra-Community acquisition (in the Member State of arrival of the goods), followed by a ‘domestic’ supply in the Member State of arrival, and requires the supplier to be identified for VAT purposes in that Member State. To avoid this, such transactions, where they take place between two taxable persons should be, under certain conditions, considered to give rise to one exempt supply in the Member State of departure and one intra-Community acquisition in the Member State of arrival.
With the completion of the single market as from the first day of 1993, fiscal controls at intra-Community frontiers were abolished. The elimination of these frontier controls induced a change in fiscal policy towards the acquisition principle.
This policy applied first to the exchanges of goods between taxable persons but was then extended with the exchanges of services between taxable persons.
The basic rule applied to the acquisition principle is that the sale (supply) of goods or services in intra-Community trade in one Member State gives rise to a corresponding acquisition (purchase) in the Member State of arrival of the goods or provision of the services. 
''')

p = document.add_paragraph('''A purchaser may take an intra-Community supply exempt of VAT as long as he meets two simple conditions:
-	He must be a taxable person;
-	Evidence must be held to prove that the goods have left the Member State of departure.
''')

p = document.add_paragraph('''Each trader doing intra-Community purchases will declare such business on his normal periodic VAT return (two global sums for the period: one for the total of his intra-Community supplies and one for the total of his intra-Community acquisitions).
Each trader making intra-Community supplies must also submit a recapitulative statement, showing the VAT identification number and net turnover value of his supplies to each of his intra-Community customers for the period. This recapitulative statement is submitted monthly, subject with the following exceptions:
•	In some Member States, the administration can allow a supplier to submit the recapitulative statement on a quarterly basis for services supplies.
•	In some Member States, the administration can, at its discretion, permit suppliers that make limited intra-community supplies to submit the recapitulative statement on a quarterly basis. Permission can be granted only if the value of the intra-community supplies in any of the current or past four quarters does not exceed a threshold. This threshold is currently €50 000.
For a supplier that is permitted to make such a quarterly declaration, it may happen that the threshold is exceeded during the quarter. In this case, the supplier must switch to monthly declarations without delay. If the threshold is exceeded in the second month of the quarter, either a recapitulative statement for the first two months of the quarter or, if required by national legislation, separate monthly recapitulative statements for these months must be submitted.
•	In some Member States, small traders are permitted to submit the recapitulative statement for a period covering more than one quarter but not exceeding one year. 
''')
document.add_heading('II.1.1.2	Triangular trade')

p = document.add_paragraph('''Triangular trade occurs between three Member States when a trader from a third Member State acts as intermediary between the seller and the final purchaser. 
Trader B is then entitled to benefit from a simplified procedure in this case of triangular trade. He does not have to be registered to the VAT in C's Member State in order to be able to acquire goods from A's country for resale to C. In addition, he does not have to pay any VAT on the amount of the invoice sent by A.
The following documents are to be filled in by each of the partners in the triangular transaction:
	Trader A :
	issues a standard invoice to B.
	declares the sale to B in his periodic VAT return and quarterly recapitulative statement.
	Trader B :
	has nothing to report for the transaction with A.
	issues an invoice to C, with a mention that the transaction relates to "triangular trade".
	declares the sale to C in his periodic VAT return.
	declares the sale to C in his recapitulative statement. This has to be identified distinctively from the other (non-triangular) transactions.
	Trader C :
	declares the purchase from B in his periodic VAT return.
''')

p = document.add_paragraph('''The purpose of the mechanism is not to allow a completely automatic reconstruction of the triangular transaction by any of the three tax administrations involved. It is rather meant to give a hint to B's administration on the apparent discrepancy, which arises: Why does B not declare the transaction with A, although B has been invoiced by A? Any further investigation can, if necessary, be performed through the follow-up request system.
The triangular trade principle does not apply to B2B services transactions.
''')
document.add_heading('II.1.1.3	Third Member State Requests')


p = document.add_paragraph('''In the context of the fight against VAT fraud, and in particular carrousel fraud, it is useful for Member States to be able to identify rapidly other customers of a given supplier in other Member States. This information is obtainable via the Third Member State requests.
A Third Member State request allows Member State A to request, for a particular supplier established in Member State B, all purchasers in all other Member States for this supplier.
The implementation of the functionality to support Third Member State requests is mandatory for all Member States.
''')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

document.add_heading('Heading, level 1', level=1)
document.add_heading('Heading, level 2', level=2)
document.add_paragraph('Intense quote', style='Intense Quote')

document.add_paragraph(
    'first item in unordered list', style='List Bullet'
)
document.add_paragraph(
    'first item in ordered list', style='List Number'
)

document.add_picture('plot.png', width=Mm(50))
document.add_picture('kep.png', width=Mm(60))
document.add_picture(buf, width=Mm(60))

records = (
    (3, '101', 'Spam'),
    (7, '422', 'Eggs'),
    (4, '631', 'Spam, spam, eggs, and spam')
)

table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
for qty, id, desc in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc

document.add_page_break()

document.save('demo.docx')

os.startfile('demo.docx')
