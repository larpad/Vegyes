import os
from pathlib import Path
from markdown import markdown
from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from bs4 import BeautifulSoup

def convert_md_to_docx(input_dir, output_dir):
    # Ensure output directory exists
    Path(output_dir).mkdir(parents=True, exist_ok=True)

    # Process all .md files in the input directory
    for md_file in Path(input_dir).glob('*.md'):
        # Read markdown content
        with open(md_file, 'r', encoding='utf-8') as file:
            md_content = file.read()

        # Convert markdown to HTML
        html = markdown(md_content, extensions=['tables'])

        # Parse HTML with BeautifulSoup
        soup = BeautifulSoup(html, 'html.parser')

        # Create a new Word document
        doc = Document()

        # Define styles
        styles = doc.styles
        style_normal = styles['Normal']
        style_normal.font.name = 'Times New Roman'
        style_normal.font.size = Pt(12)

        # Process the HTML content
        for element in soup.find_all(['p', 'h1', 'h2', 'h3', 'table']):
            if element.name == 'table':
                # Convert HTML table to Word table
                table = doc.add_table(rows=0, cols=len(element.find_all('th')))
                for row in element.find_all('tr'):
                    cells = row.find_all(['th', 'td'])
                    word_row = table.add_row().cells
                    for i, cell in enumerate(cells):
                        word_row[i].text = cell.get_text(strip=True)
                        if cell.name == 'th':
                            word_row[i].paragraphs[0].runs[0].bold = True
            else:
                # Add other elements as paragraphs
                p = doc.add_paragraph(element.get_text())
                if element.name.startswith('h'):
                    p.style = f'Heading {element.name[1]}'

        # Save the document
        output_file = Path(output_dir) / f"{md_file.stem}.docx"
        doc.save(output_file)
        print(f"Converted {md_file} to {output_file}")

if __name__ == "__main__":
    input_directory = "."  # Current directory
    output_directory = "./converted_docs"
    convert_md_to_docx(input_directory, output_directory)
